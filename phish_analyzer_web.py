#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Mansimov Phishing Email Analyzer - Web Version (Streamlit)
pip install streamlit matplotlib python-docx eml-parser extract-msg oletools yara-python pefile beautifulsoup4 lxml tldextract chardet aiohttp aiodns requests[security] python-whois dnspython pyyaml rich jinja2 python-dateutil
streamlit run phish_analyzer_web.py

(typically http://localhost:8501)

- Load .eml/.msg
- Analyze (async) with URL expansion, YARA, enrichment (VT, OTX, urlscan, GN, AbuseIPDB, WHOIS)
- Visualize IOCs, score, ATT&CK tags, attachments
- Export Word (.docx) report with embedded charts

Safe-by-design: no dynamic execution of attachments; only parsing + HTTP lookups with TLS.

Developed by N.Mansimov
"""
import streamlit as st
import asyncio
import tempfile
import os
import re
import json
import base64
from datetime import datetime
from pathlib import Path
from typing import Dict, Any, List, Optional, Tuple

# Core Analysis Imports
import argparse
import binascii
import email
import email.policy
import hashlib
import io
import zipfile
import aiohttp
import chardet
import jinja2
import pefile
import tldextract
import whois
import yaml
import yara
from bs4 import BeautifulSoup
from email.header import decode_header, make_header
from email.utils import parsedate_to_datetime
from dateutil import tz

# Optional imports guarded (graceful degrade)
try:
    import extract_msg
except Exception:
    extract_msg = None

try:
    from oletools.olevba import VBA_Parser, TYPE_OLE, TYPE_OpenXML, TYPE_Word2003_XML, TYPE_MHTML
except Exception:
    VBA_Parser = None

# For charts
import matplotlib.pyplot as plt
import matplotlib
matplotlib.use('Agg')  # Use non-interactive backend

# Word export
from docx import Document
from docx.shared import Inches

# Constants
URL_GENERIC_RE = re.compile(
    r'(?:(?:hxxp|http|https):\/\/)?[a-zA-Z0-9\-\.\[\]]+\.[a-zA-Z]{2,}(?:[:\/][^\s<>"\)\]]*)?',
    re.IGNORECASE
)
IP_RE = re.compile(r'\b(?:\d{1,3}\.){3}\d{1,3}\b')
HASH_RE = re.compile(r'\b[a-fA-F0-9]{32}\b|\b[a-fA-F0-9]{40}\b|\b[a-fA-F0-9]{64}\b')

SHORTENERS = {"bit.ly","goo.gl","t.co","tinyurl.com","ow.ly","is.gd","buff.ly","cutt.ly","lnkd.in","rebrand.ly"}
SUS_TLDS = {"zip","mov","country","gq","work","click","top","xyz","support","help","pics"}

ATTACK_TECHNIQUES = {
    "T1566": "Phishing",
    "T1566.001": "Spearphishing Attachment",
    "T1566.002": "Spearphishing Link",
    "T1204": "User Execution",
    "T1140": "Deobfuscate/Decode",
    "T1059": "Command and Scripting Interpreter",
    "T1105": "Ingress Tool Transfer",
}

DEFAULT_CFG: Dict[str, Any] = {
    "apis": {
        "virustotal_api_key": "06e809ecf15d400ac95f5376765eff953e4225c0846d83a795fcb1b669d674c2",
        "otx_api_key": "ffa4f89b675c44f5f4d04b7b823720212095ef37c300c0564ffe81c78546df48",
        "urlscan_api_key": "0198e4cd-9856-71fe-a91d-e4f249908c2e",
        "greynoise_api_key": "fa909904-bf79-4cad-8f84-a43aa2659632",
        "abuseipdb_api_key": "fcab3a321d5debbab5b6f8fd32bf7847a9de868935f42daae14403a951c5490f3e6c7bd3ed630a3c"
    },
    "network": {
        "user_agent": "Mansimov-Phish-Analyzer/1.0",
        "timeout_sec": 10,
        "verify_tls": True,
        "safe_http_methods": ["HEAD", "GET"],
        "expand_max_redirects": 5
    },
    "scoring": {
        "weights": {
            "url_new_domain": 10,
            "url_sus_tld": 8,
            "url_ip_host": 12,
            "url_shortener": 6,
            "macro_autoexec": 20,
            "attachment_exe": 25,
            "attachment_script": 15,
            "vt_malicious_detection": 30,
            "brand_impersonation": 12,
            "dmarc_fail": 10,
            "replyto_mismatch": 8,
            "unicode_spoof": 10,
            "otx_pulse_hit": 10,
            "greynoise_malicious": 8,
            "abuseipdb_reports": 8
        }
    },
    "output": {
        "html_report": True,
        "markdown_report": True,
        "json_report": True
    },
    "yara": {
        "rules_path": "yara_rules"
    }
}

# Analysis functions (same as desktop version)
def load_config(path: Optional[str]) -> Dict[str, Any]:
    if not path:
        return DEFAULT_CFG
    p = Path(path)
    if not p.exists():
        return DEFAULT_CFG
    with open(p, "r", encoding="utf-8") as f:
        d = yaml.safe_load(f) or {}
    # Merge into defaults (shallow)
    cfg = json.loads(json.dumps(DEFAULT_CFG))
    for k, v in d.items():
        cfg[k] = v
    return cfg

def safe_decode(bytes_or_str: Any) -> str:
    if isinstance(bytes_or_str, str):
        return bytes_or_str
    if bytes_or_str is None:
        return ""
    guess = chardet.detect(bytes_or_str)
    enc = guess.get("encoding") or "utf-8"
    try:
        return bytes_or_str.decode(enc, errors="replace")
    except Exception:
        return bytes_or_str.decode("utf-8", errors="replace")

def normalize_url(u: str) -> str:
    s = u.strip()
    s = s.replace("[.]", ".").replace("[DOT]", ".").replace("[dot]", ".")
    s = s.replace("hxxp://", "http://").replace("hxxps://", "https://")
    s = re.sub(r'\s+', '', s)
    return s

def parse_eml(path: Path) -> Dict[str, Any]:
    with open(path, "rb") as f:
        msg = email.message_from_binary_file(f, policy=email.policy.default)
    return message_to_dict(msg)

def parse_msg(path: Path) -> Dict[str, Any]:
    if extract_msg is None:
        raise RuntimeError("extract_msg not installed")
    msg = extract_msg.Message(str(path))
    headers = {k: v for k, v in msg.headerItems()}
    dt = None
    try:
        dt = parsedate_to_datetime(headers.get("Date")) if headers.get("Date") else None
    except Exception:
        dt = None
    attachments = []
    for att in msg.attachments:
        attachments.append({
            "filename": att.longFilename or att.shortFilename,
            "content": att.data,
            "content_type": "application/octet-stream"
        })
    return {
        "headers": headers,
        "date": dt.isoformat() if dt else None,
        "subject": msg.subject or "",
        "from": headers.get("From", ""),
        "to": headers.get("To", ""),
        "reply_to": headers.get("Reply-To", ""),
        "body_text": safe_decode(msg.body.encode("utf-8")) if msg.body else "",
        "body_html": safe_decode((msg.htmlBody or "").encode("utf-8")),
        "attachments": attachments,
        "raw_source_path": str(path)
    }

def message_to_dict(msg: email.message.EmailMessage) -> Dict[str, Any]:
    headers = {k: str(make_header(decode_header(v))) for k, v in msg.items()}
    dt = None
    try:
        dt = parsedate_to_datetime(headers.get("Date")) if headers.get("Date") else None
    except Exception:
        dt = None
    body_text, body_html = "", ""
    attachments = []
    if msg.is_multipart():
        for part in msg.walk():
            ctype = part.get_content_type() or ""
            disp = (part.get("Content-Disposition") or "").lower()
            if ctype == "text/plain" and "attachment" not in disp:
                body_text += safe_decode(part.get_payload(decode=True))
            elif ctype == "text/html" and "attachment" not in disp:
                body_html += safe_decode(part.get_payload(decode=True))
            elif "attachment" in disp or part.get_filename():
                attachments.append({
                    "filename": part.get_filename(),
                    "content_type": ctype,
                    "content": part.get_payload(decode=True) or b""
                })
    else:
        ctype = msg.get_content_type()
        payload = msg.get_payload(decode=True) or b""
        if ctype == "text/plain":
            body_text = safe_decode(payload)
        elif ctype == "text/html":
            body_html = safe_decode(payload)
    return {
        "headers": headers,
        "date": dt.isoformat() if dt else None,
        "subject": headers.get("Subject", ""),
        "from": headers.get("From", ""),
        "to": headers.get("To", ""),
        "reply_to": headers.get("Reply-To", ""),
        "body_text": body_text,
        "body_html": body_html,
        "attachments": attachments,
        "raw_source_path": None
    }

def extract_urls(text: str) -> List[str]:
    urls = set()
    for m in URL_GENERIC_RE.finditer(text):
        urls.add(normalize_url(m.group(0)))
    return sorted(u for u in urls if "." in u)

def extract_iocs(email_obj: Dict[str, Any]) -> Dict[str, List[str]]:
    texts = " ".join([
        email_obj.get("subject",""),
        email_obj.get("from",""),
        email_obj.get("to",""),
        email_obj.get("reply_to",""),
        email_obj.get("body_text",""),
        email_obj.get("body_html","")
    ])
    urls = set(extract_urls(texts))
    ips = set(IP_RE.findall(texts))
    hashes = set(HASH_RE.findall(texts))
    soup = BeautifulSoup(email_obj.get("body_html",""), "lxml")
    for tag in soup.find_all(["a","img","form","script"]):
        for attr in ["href","src","action"]:
            val = tag.get(attr)
            if val:
                urls.add(normalize_url(val))
    domains = set()
    for u in list(urls):
        if "://" not in u:
            u = "http://" + u
        try:
            ext = tldextract.extract(u)
            if ext.domain and ext.suffix:
                domains.add(".".join(p for p in [ext.subdomain, ext.domain, ext.suffix] if p))
        except Exception:
            pass
    return {
        "urls": sorted(urls),
        "ips": sorted(ips),
        "domains": sorted(domains),
        "hashes": sorted(hashes)
    }

def header_auth_findings(headers: Dict[str,str]) -> List[str]:
    f = []
    spf = headers.get("Received-SPF","")
    authres = headers.get("Authentication-Results","")
    if "fail" in spf.lower():
        f.append("SPF fail observed in Received-SPF")
    if "dmarc=fail" in authres.lower():
        f.append("DMARC fail in Authentication-Results")
    if "dkim=fail" in authres.lower():
        f.append("DKIM fail in Authentication-Results")
    rp = headers.get("Return-Path","")
    frm = headers.get("From","")
    if rp and frm and rp.split("@")[-1].strip("<> ") != frm.split("@")[-1].strip("<> "):
        f.append("Return-Path domain mismatch with From")
    return f

def brand_impersonation_indicators(subject: str, body_html: str) -> List[str]:
    brands = ["Microsoft", "Office 365", "PayPal", "DocuSign", "Adobe", "Amazon", "OneDrive", "Dropbox", "Okta", "GitHub", "Google"]
    hits = []
    for b in brands:
        if re.search(rf'\b{re.escape(b)}\b', subject, re.I) or re.search(rf'\b{re.escape(b)}\b', body_html, re.I):
            hits.append(f"Brand keyword present: {b}")
    return hits

def contains_unicode_spoofers(s: str) -> bool:
    if re.search(r'[\u200B-\u200D\uFEFF]', s):
        return True
    return bool(re.search(r'[\u0400-\u04FF]', s))

def hash_bytes(b: bytes) -> Dict[str,str]:
    return {
        "md5": hashlib.md5(b).hexdigest(),
        "sha1": hashlib.sha1(b).hexdigest(),
        "sha256": hashlib.sha256(b).hexdigest()
    }

def analyze_attachments(atts: List[Dict[str,Any]], yara_rules: Optional[yara.Rules]) -> List[Dict[str,Any]]:
    results = []
    for att in atts:
        name = att.get("filename") or "unknown"
        ctype = att.get("content_type") or "application/octet-stream"
        content: bytes = att.get("content") or b""
        h = hash_bytes(content)
        finding = {
            "filename": name,
            "content_type": ctype,
            "size": len(content),
            "hashes": h,
            "macro_analysis": None,
            "pe_analysis": None,
            "yara_matches": [],
            "pe": False,
            "office_doc": False,
            "script_like": False
        }
        if yara_rules:
            try:
                m = yara_rules.match(data=content, timeout=3.0)
                finding["yara_matches"] = [str(i) for i in m]
            except Exception:
                pass
        if VBA_Parser is not None:
            try:
                vba = VBA_Parser(filename=name, data=content)
                finding["office_doc"] = vba.detect_vba_macros()
                if finding["office_doc"]:
                    autoexec = False
                    suspicious = []
                    for (_fname, _stream_path, _vba_filename, vba_code) in vba.extract_macros():
                        if not vba_code:
                            continue
                        if re.search(r'Auto(?:Open|_?Open|Close|_?Close|Workbook_Open|Document_Open)', vba_code, re.I):
                            autoexec = True
                        for pat in [r'CreateObject\("WScript\.Shell"\)', r'Shell\(', r'GetObject\(', r'Execute\(', r'Bitstransfer', r'URLDownloadToFile', r'WinHttp', r'MSXML2\.XMLHTTP']:
                            if re.search(pat, vba_code, re.I):
                                suspicious.append(pat)
                    finding["macro_analysis"] = {
                        "autoexec": autoexec,
                        "suspicious_patterns": list(sorted(set(suspicious)))
                    }
            except Exception:
                pass
        try:
            pe = pefile.PE(data=content, fast_load=True)
            finding["pe"] = True
            finding["pe_analysis"] = {
                "is_dll": bool(pe.FILE_HEADER.Characteristics & 0x2000),
                "machine": hex(pe.FILE_HEADER.Machine),
                "number_of_sections": pe.FILE_HEADER.NumberOfSections
            }
        except Exception:
            pass
        if name.lower().endswith((".vbs",".js",".jse",".wsf",".ps1",".cmd",".bat",".lnk",".hta",".vbe")):
            finding["script_like"] = True
        results.append(finding)
    return results

async def expand_url(session: aiohttp.ClientSession, url: str, max_redirects: int) -> Dict[str, Any]:
    url_norm = normalize_url(url)
    if not url_norm.startswith("http"):
        url_norm = "http://" + url_norm
    res = {"input": url, "expanded": url_norm, "chain": [], "status": None}
    try:
        async with session.get(url_norm, allow_redirects=False) as r:
            res["status"] = r.status
            res["chain"].append({"url": url_norm, "status": r.status})
            redirects = 0
            next_url = r.headers.get("Location")
            while next_url and redirects < max_redirects:
                if not next_url.startswith("http"):
                    from yarl import URL as YURL
                    next_url = str(YURL(url_norm).join(YURL(next_url)))
                res["chain"].append({"url": next_url, "status": None})
                async with session.get(next_url, allow_redirects=False) as nr:
                    res["chain"][-1]["status"] = nr.status
                    url_norm = next_url
                    next_url = nr.headers.get("Location")
                    redirects += 1
            res["expanded"] = url_norm
    except Exception as e:
        res["error"] = str(e)
    return res

async def enrich_iocs(cfg: Dict[str,Any], iocs: Dict[str, List[str]]) -> Dict[str,Any]:
    out = {"virustotal": {}, "otx": {}, "urlscan": {}, "greynoise": {}, "abuseipdb": {}, "whois": {}}
    headers = {"User-Agent": cfg["network"]["user_agent"]}
    timeout = aiohttp.ClientTimeout(total=cfg["network"]["timeout_sec"])
    ssl = cfg["network"]["verify_tls"]
    async with aiohttp.ClientSession(timeout=timeout, headers=headers) as session:
        vt_key = cfg["apis"].get("virustotal_api_key")
        if vt_key:
            for h in iocs["hashes"][:25]:
                url = f"https://www.virustotal.com/api/v3/files/{h}"
                try:
                    async with session.get(url, headers={"x-apikey": vt_key}, ssl=ssl) as r:
                        if r.status == 200:
                            out["virustotal"][h] = await r.json()
                except Exception:
                    pass
            for d in iocs["domains"][:25]:
                url = f"https://www.virustotal.com/api/v3/domains/{d}"
                try:
                    async with session.get(url, headers={"x-apikey": vt_key}, ssl=ssl) as r:
                        if r.status == 200:
                            out["virustotal"][d] = await r.json()
                except Exception:
                    pass
        us_key = cfg["apis"].get("urlscan_api_key")
        if us_key:
            for u in iocs["urls"][:25]:
                try:
                    q = f"https://urlscan.io/api/v1/search/?q=domain:{u}"
                    headers2 = {"API-Key": us_key, "Content-Type": "application/json"}
                    async with session.get(q, headers=headers2, ssl=ssl) as r:
                        if r.status == 200:
                            out["urlscan"][u] = await r.json()
                except Exception:
                    pass
        otx_key = cfg["apis"].get("otx_api_key")
        if otx_key:
            for d in iocs["domains"][:25]:
                try:
                    u = f"https://otx.alienvault.com/api/v1/indicators/domain/{d}/general"
                    async with session.get(u, headers={"X-OTX-API-KEY": otx_key}, ssl=ssl) as r:
                        if r.status == 200:
                            out["otx"][d] = await r.json()
                except Exception:
                    pass
        gn_key = cfg["apis"].get("greynoise_api_key")
        if gn_key:
            for ip in iocs["ips"][:50]:
                try:
                    u = f"https://api.greynoise.io/v3/community/{ip}"
                    async with session.get(u, headers={"key": gn_key}, ssl=ssl) as r:
                        if r.status == 200:
                            out["greynoise"][ip] = await r.json()
                except Exception:
                    pass
        ab_key = cfg["apis"].get("abuseipdb_api_key")
        if ab_key:
            for ip in iocs["ips"][:50]:
                try:
                    u = f"https://api.abuseipdb.com/api/v2/check?ipAddress={ip}&maxAgeInDays=90"
                    async with session.get(u, headers={"Key": ab_key, "Accept": "application/json"}, ssl=ssl) as r:
                        if r.status == 200:
                            out["abuseipdb"][ip] = await r.json()
                except Exception:
                    pass
        for d in iocs["domains"][:25]:
            try:
                w = whois.whois(d)
                out["whois"][d] = {
                    "creation_date": str(w.creation_date),
                    "updated_date": str(w.updated_date),
                    "registrar": w.registrar,
                    "name_servers": w.name_servers
                }
            except Exception:
                pass
    return out

def score_findings(cfg: Dict[str,Any], email_obj: Dict[str,Any], iocs: Dict[str,List[str]], att_results: List[Dict[str,Any]], enrich: Dict[str,Any]) -> Tuple[int, List[str], List[str]]:
    w = cfg["scoring"]["weights"]
    score = 0
    tags = set()
    notes = []

    for u in iocs["urls"]:
        try:
            ext = tldextract.extract(u if u.startswith("http") else "http://" + u)
            if ext and ext.suffix:
                if ext.suffix in SUS_TLDS:
                    score += w["url_sus_tld"]; notes.append(f"suspicious TLD: .{ext.suffix}")
        except Exception:
            pass
        if re.match(r'^https?://\d{1,3}(?:\.\d{1,3}){3}', u):
            score += w["url_ip_host"]; notes.append("URL host is raw IP")
        host = (tldextract.extract(u).fqdn or "").lower()
        if any(host.endswith(s) for s in SHORTENERS):
            score += w["url_shortener"]; notes.append("URL shortener used")
        tags.add("T1566.002")

    for a in att_results:
        if a["macro_analysis"]:
            if a["macro_analysis"]["autoexec"]:
                score += w["macro_autoexec"]; notes.append(f"Macro autoexec in {a['filename']}")
            if a["macro_analysis"]["suspicious_patterns"]:
                score += 8; notes.append(f"Suspicious VBA in {a['filename']}")
            tags.add("T1566.001"); tags.add("T1204")
        if a["pe"]:
            score += w["attachment_exe"]; notes.append(f"Executable attachment: {a['filename']}")
            tags.add("T1204")
        if a["script_like"]:
            score += w["attachment_script"]; notes.append(f"Script-like attachment: {a['filename']}")
        if a["yara_matches"]:
            score += 10; notes.append(f"YARA hits: {', '.join(a['yara_matches'])}")

    for h in header_auth_findings(email_obj["headers"]):
        notes.append(h)
    if any("DMARC fail" in n for n in notes):
        score += w["dmarc_fail"]
    if "Return-Path domain mismatch with From" in notes:
        score += w["replyto_mismatch"]

    brand_hits = brand_impersonation_indicators(email_obj.get("subject",""), email_obj.get("body_html",""))
    if brand_hits:
        notes.extend(brand_hits)
        score += w["brand_impersonation"]

    if contains_unicode_spoofers(email_obj.get("subject","") + email_obj.get("body_html","")):
        score += w["unicode_spoof"]; notes.append("Unicode zero-width/homoglyphs present")

    vt = enrich.get("virustotal", {})
    for key, val in vt.items():
        try:
            stats = val["data"]["attributes"]["last_analysis_stats"]
            if stats.get("malicious", 0) >= 3:
                score += w["vt_malicious_detection"]
                notes.append(f"VirusTotal indicates malicious for {key} ({stats.get('malicious')} vendors)")
        except Exception:
            pass

    for ip, v in enrich.get("abuseipdb", {}).items():
        try:
            rep = v["data"]["totalReports"]
            if rep and rep >= 5:
                score += w["abuseipdb_reports"]; notes.append(f"AbuseIPDB {ip} reports: {rep}")
        except Exception:
            pass

    for ip, v in enrich.get("greynoise", {}).items():
        try:
            if v.get("classification") == "malicious":
                score += w["greynoise_malicious"]; notes.append(f"GreyNoise malicious IP: {ip}")
        except Exception:
            pass

    for d, v in enrich.get("otx", {}).items():
        try:
            pulses = v.get("pulse_info", {}).get("count", 0)
            if pulses > 0:
                score += w["otx_pulse_hit"]; notes.append(f"OTX pulses hit: {d} ({pulses})")
        except Exception:
            pass

    if any("Macro" in n for n in notes) or any(a["macro_analysis"] for a in att_results):
        tags.add("T1140")
    if any(("URL shortener" in n) or ("URL host is raw IP" in n) for n in notes):
        tags.add("T1105")

    return score, sorted(tags), notes

def render_markdown(email_obj: Dict[str,Any], iocs: Dict[str,List[str]], att_results: List[Dict[str,Any]], score: int, tags: List[str], notes: List[str]) -> str:
    md = []
    md.append(f"# Phishing Analysis Report\n")
    md.append(f"- Subject: {email_obj.get('subject','')}")
    md.append(f"- From: {email_obj.get('from','')}")
    md.append(f"- To: {email_obj.get('to','')}")
    md.append(f"- Date: {email_obj.get('date')}")
    md.append(f"- Score: {score}")
    items = []
    for t in tags:
        technique = ATTACK_TECHNIQUES.get(t, "?")
        items.append(f"{t} ({technique})")
    
    line = "- MITRE ATT&CK: " + ", ".join(items)
    md.append(line)
    md.append("\n## Findings")
    for n in notes:
        md.append(f"- {n}")
    md.append("\n## IOCs")
    md.append(f"- Domains: {', '.join(iocs['domains'])}")
    md.append(f"- IPs: {', '.join(iocs['ips'])}")
    md.append(f"- URLs:")
    for u in iocs["urls"]:
        md.append(f"  - {u}")
    md.append(f"- Hashes: {', '.join(iocs['hashes'])}")
    md.append("\n## Attachments")
    for a in att_results:
        md.append(f"- {a['filename']} ({a['content_type']}, {a['size']} bytes)")
        md.append(f"  - Hashes: md5={a['hashes']['md5']} sha256={a['hashes']['sha256']}")
        if a["macro_analysis"]:
            md.append(f"  - Macro autoexec: {a['macro_analysis']['autoexec']}")
            if a["macro_analysis"]["suspicious_patterns"]:
                md.append(f"  - Suspicious VBA: {', '.join(a['macro_analysis']['suspicious_patterns'])}")
        if a["pe"]:
            md.append(f"  - PE: {json.dumps(a['pe_analysis'])}")
        if a["yara_matches"]:
            md.append(f"  - YARA: {', '.join(a['yara_matches'])}")
    return "\n".join(md)

async def analyze_email_async(
    email_path: str,
    cfg_path: Optional[str],
    yara_dir: Optional[str],
    expand: bool
) -> Dict[str, Any]:
    cfg = load_config(cfg_path)
    p = Path(email_path)
    if p.suffix.lower() == ".eml":
        email_obj = parse_eml(p)
    elif p.suffix.lower() == ".msg":
        email_obj = parse_msg(p)
    else:
        raise RuntimeError("Unsupported file type. Use .eml or .msg")

    yara_rules = None
    if yara_dir:
        yr = Path(yara_dir)
        if yr.exists():
            rule_files = [str(r) for r in yr.rglob("*.yar")]
            if rule_files:
                filemap = {f"r{i}": rf for i, rf in enumerate(rule_files)}
                try:
                    yara_rules = yara.compile(filepaths=filemap)
                except Exception as e:
                    # proceed without YARA
                    yara_rules = None

    att_results = analyze_attachments(email_obj["attachments"], yara_rules)
    iocs = extract_iocs(email_obj)

    expansions = []
    if expand and iocs["urls"]:
        timeout = aiohttp.ClientTimeout(total=cfg["network"]["timeout_sec"])
        async with aiohttp.ClientSession(timeout=timeout, headers={"User-Agent": cfg["network"]["user_agent"]}) as session:
            tasks = [expand_url(session, u, cfg["network"]["expand_max_redirects"]) for u in iocs["urls"][:30]]
            expansions = await asyncio.gather(*tasks)
        expanded = [e.get("expanded", e["input"]) for e in expansions]
        iocs["urls"] = sorted(set(expanded + iocs["urls"]))

    enrichment = await enrich_iocs(cfg, iocs)
    score, tags, notes = score_findings(cfg, email_obj, iocs, att_results, enrichment)
    md = render_markdown(email_obj, iocs, att_results, score, tags, notes)

    return {
        "email": email_obj,
        "iocs": iocs,
        "attachments": att_results,
        "score": score,
        "tags": tags,
        "notes": notes,
        "enrichment": enrichment,
        "expansions": expansions,
        "markdown": md,
        "config_used": cfg
    }

# Streamlit UI
def main():
    st.set_page_config(
        page_title="Phishing Email Analyzer",
        page_icon="📧",
        layout="wide",
        initial_sidebar_state="expanded"
    )
    
    st.title("📧 Phishing Email Analyzer")
    st.markdown("Analyze .eml and .msg files for phishing indicators")
    
    # Initialize session state
    if "analysis_result" not in st.session_state:
        st.session_state.analysis_result = None
    if "analysis_error" not in st.session_state:
        st.session_state.analysis_error = None
    
    # Sidebar for inputs
    with st.sidebar:
        st.header("Input Settings")
        
        uploaded_file = st.file_uploader("Upload Email File", type=["eml", "msg"])
        
        expand_urls = st.checkbox("Expand URLs", value=True)
        
        yara_dir = st.text_input("YARA Rules Directory (optional)", "")
        
        config_file = st.file_uploader("Upload Config File (optional)", type=["yaml", "yml"])
        
        analyze_button = st.button("Analyze Email", type="primary")
    
    # Main content area
    if analyze_button and uploaded_file:
        with st.spinner("Analyzing email..."):
            # Save uploaded files to temporary location
            with tempfile.NamedTemporaryFile(delete=False, suffix=f".{uploaded_file.name.split('.')[-1]}") as tmp_file:
                tmp_file.write(uploaded_file.getvalue())
                email_path = tmp_file.name
            
            config_path = None
            if config_file:
                with tempfile.NamedTemporaryFile(delete=False, suffix=".yaml") as cfg_file:
                    cfg_file.write(config_file.getvalue())
                    config_path = cfg_file.name
            
            try:
                # Run analysis
                result = asyncio.run(analyze_email_async(
                    email_path, config_path, yara_dir, expand_urls
                ))
                st.session_state.analysis_result = result
                st.session_state.analysis_error = None
            except Exception as e:
                st.session_state.analysis_result = None
                st.session_state.analysis_error = str(e)
            finally:
                # Clean up temp files
                os.unlink(email_path)
                if config_path:
                    os.unlink(config_path)
    
    # Display results
    if st.session_state.analysis_error:
        st.error(f"Analysis failed: {st.session_state.analysis_error}")
    
    if st.session_state.analysis_result:
        result = st.session_state.analysis_result
        email_obj = result["email"]
        iocs = result["iocs"]
        atts = result["attachments"]
        score = result["score"]
        tags = result["tags"]
        notes = result["notes"]
        
        # Summary section
        st.header("Analysis Summary")
        col1, col2, col3 = st.columns(3)
        
        with col1:
            st.metric("Phishing Score", score)
        
        with col2:
            items = []
            for t in tags:
                technique = ATTACK_TECHNIQUES.get(t, "?")
                items.append(f"{t} ({technique})")
            st.write("MITRE ATT&CK:", ", ".join(items))
        
        with col3:
            st.write("Subject:", email_obj.get("subject", ""))
            st.write("From:", email_obj.get("from", ""))
            st.write("Date:", email_obj.get("date", ""))
        
        # Create charts
        st.subheader("Dashboard")
        fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(10, 4))
        
        # IOC counts chart
        ioc_counts = {
            "Domains": len(iocs["domains"]),
            "IPs": len(iocs["ips"]),
            "URLs": len(iocs["urls"]),
            "Hashes": len(iocs["hashes"]),
        }
        ax1.bar(ioc_counts.keys(), ioc_counts.values(), color=["#2b8a3e","#1971c2","#e8590c","#862e9c"])
        ax1.set_title("IOC Counts")
        ax1.set_ylabel("Count")
        ax1.tick_params(axis='x', rotation=20)
        
        # Score donut chart
        display_score = max(0, min(100, score))
        sizes = [display_score, 100 - display_score]
        colors = ["#d9480f", "#e9ecef"] if display_score >= 60 else ("#f08c00" if display_score >= 30 else "#2f9e44", "#e9ecef")
        wedges, _ = ax2.pie(sizes, colors=colors, startangle=90, counterclock=False, wedgeprops=dict(width=0.35))
        ax2.set_aspect('equal')
        ax2.set_title(f"Score: {score}")
        
        st.pyplot(fig)
        
        # Findings section
        st.subheader("Findings")
        for note in notes:
            st.write(f"- {note}")
        
        # Tabs for detailed information
        tab1, tab2, tab3, tab4 = st.tabs(["IOCs", "Attachments", "Enrichment", "Full Report"])
        
        with tab1:
            st.subheader("Indicators of Compromise")
            
            col1, col2 = st.columns(2)
            
            with col1:
                st.write("**Domains**")
                for domain in iocs["domains"]:
                    st.code(domain)
            
            with col2:
                st.write("**IP Addresses**")
                for ip in iocs["ips"]:
                    st.code(ip)
            
            st.write("**URLs**")
            for url in iocs["urls"]:
                st.code(url)
            
            st.write("**Hashes**")
            for hash_val in iocs["hashes"]:
                st.code(hash_val)
        
        with tab2:
            st.subheader("Attachment Analysis")
            if atts:
                for att in atts:
                    with st.expander(att["filename"]):
                        st.write(f"Type: {att['content_type']}")
                        st.write(f"Size: {att['size']} bytes")
                        st.write(f"MD5: {att['hashes']['md5']}")
                        st.write(f"SHA256: {att['hashes']['sha256']}")
                        
                        if att["macro_analysis"]:
                            st.write("Macro Analysis:")
                            st.write(f"- AutoExec: {att['macro_analysis']['autoexec']}")
                            if att["macro_analysis"]["suspicious_patterns"]:
                                st.write("- Suspicious Patterns:")
                                for pattern in att["macro_analysis"]["suspicious_patterns"]:
                                    st.write(f"  - {pattern}")
                        
                        if att["pe"]:
                            st.write("PE File Analysis:")
                            st.json(att["pe_analysis"])
                        
                        if att["yara_matches"]:
                            st.write("YARA Matches:")
                            for match in att["yara_matches"]:
                                st.write(f"- {match}")
            else:
                st.write("No attachments found")
        
        with tab3:
            st.subheader("Enrichment Data")
            enrichment = result["enrichment"]
            
            # Display a subset of enrichment data
            if enrichment["virustotal"]:
                st.write("VirusTotal Results")
                st.json(list(enrichment["virustotal"].keys())[:5])
            
            if enrichment["otx"]:
                st.write("AlienVault OTX Results")
                st.json(list(enrichment["otx"].keys())[:5])
            
            if enrichment["abuseipdb"]:
                st.write("AbuseIPDB Results")
                st.json(list(enrichment["abuseipdb"].keys())[:5])
        
        with tab4:
            st.subheader("Full Report")
            st.download_button(
                label="Download Word Report",
                data=generate_word_report(result),
                file_name=f"phish_report_{datetime.utcnow().strftime('%Y%m%d_%H%M%S')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            st.text_area("Markdown Report", result["markdown"], height=400)
    
    elif not st.session_state.analysis_result:
        st.info("Upload an email file and click 'Analyze Email' to get started.")

def generate_word_report(result: Dict[str, Any]) -> bytes:
    """Generate a Word document report from analysis results"""
    email_obj = result["email"]
    iocs = result["iocs"]
    atts = result["attachments"]
    score = result["score"]
    tags = result["tags"]
    notes = result["notes"]
    
    # Create a document
    doc = Document()
    
    # Add title
    doc.add_heading('Phishing Analysis Report', 0)
    
    # Add summary information
    doc.add_heading('Summary', level=1)
    p = doc.add_paragraph()
    p.add_run('Subject: ').bold = True
    p.add_run(email_obj.get('subject', ''))
    
    p = doc.add_paragraph()
    p.add_run('From: ').bold = True
    p.add_run(email_obj.get('from', ''))
    
    p = doc.add_paragraph()
    p.add_run('To: ').bold = True
    p.add_run(email_obj.get('to', ''))
    
    p = doc.add_paragraph()
    p.add_run('Date: ').bold = True
    p.add_run(str(email_obj.get('date', '')))
    
    p = doc.add_paragraph()
    p.add_run('Score: ').bold = True
    p.add_run(str(score))
    
    p = doc.add_paragraph()
    p.add_run('MITRE ATT&CK: ').bold = True
    items = []
    for t in tags:
        technique = ATTACK_TECHNIQUES.get(t, "?")
        items.append(f"{t} ({technique})")
    p.add_run(", ".join(items))
    
    # Add findings
    doc.add_heading('Findings', level=1)
    for note in notes:
        doc.add_paragraph(note, style='List Bullet')
    
    # Add IOCs
    doc.add_heading('Indicators of Compromise', level=1)
    
    doc.add_heading('Domains', level=2)
    for domain in iocs["domains"]:
        doc.add_paragraph(domain, style='List Bullet')
    
    doc.add_heading('IP Addresses', level=2)
    for ip in iocs["ips"]:
        doc.add_paragraph(ip, style='List Bullet')
    
    doc.add_heading('URLs', level=2)
    for url in iocs["urls"]:
        doc.add_paragraph(url, style='List Bullet')
    
    doc.add_heading('Hashes', level=2)
    for hash_val in iocs["hashes"]:
        doc.add_paragraph(hash_val, style='List Bullet')
    
    # Add attachments
    if atts:
        doc.add_heading('Attachments', level=1)
        for att in atts:
            doc.add_heading(att['filename'], level=2)
            doc.add_paragraph(f"Type: {att['content_type']}")
            doc.add_paragraph(f"Size: {att['size']} bytes")
            doc.add_paragraph(f"MD5: {att['hashes']['md5']}")
            doc.add_paragraph(f"SHA256: {att['hashes']['sha256']}")
            
            if att["macro_analysis"]:
                doc.add_paragraph("Macro Analysis:")
                doc.add_paragraph(f"AutoExec: {att['macro_analysis']['autoexec']}", style='List Bullet')
                if att["macro_analysis"]["suspicious_patterns"]:
                    doc.add_paragraph("Suspicious Patterns:")
                    for pattern in att["macro_analysis"]["suspicious_patterns"]:
                        doc.add_paragraph(pattern, style='List Bullet 2')
    
    # Save document to a bytes buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer.getvalue()

if __name__ == "__main__":
    main()
