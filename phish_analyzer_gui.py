#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Mansimov Phishing Email Analyzer - PyQt5 Desktop GUI
pip install PyQt5 matplotlib python-docx eml-parser extract-msg oletools yara-python pefile beautifulsoup4 lxml tldextract chardet aiohttp aiodns requests[security] python-whois dnspython pyyaml rich jinja2 python-dateutil

- Load .eml/.msg
- Analyze (async) with URL expansion, YARA, enrichment (VT, OTX, urlscan, GN, AbuseIPDB, WHOIS)
- Visualize IOCs, score, ATT&CK tags, attachments
- Export Word (.docx) report with embedded charts

Safe-by-design: no dynamic execution of attachments; only parsing + HTTP lookups with TLS.

Developed by N.Mansimov
"""
import sys
import os
import re
import json
import asyncio
from datetime import datetime
from pathlib import Path
from typing import Dict, Any, List, Optional, Tuple

# ---------- Core Analysis Imports (from your engine) ----------
import argparse
import base64
import binascii
import email
import email.policy
import hashlib
import io
import zipfile
import aiohttp
import chardet
import jinja2
import pefile
import tldextract
import whois
import yaml
import yara
from bs4 import BeautifulSoup
from email.header import decode_header, make_header
from email.utils import parsedate_to_datetime
from dateutil import tz

# Optional imports guarded (graceful degrade)
try:
    import extract_msg
except Exception:
    extract_msg = None

try:
    from oletools.olevba import VBA_Parser, TYPE_OLE, TYPE_OpenXML, TYPE_Word2003_XML, TYPE_MHTML
except Exception:
    VBA_Parser = None

# ---------- GUI & Charts ----------
from PyQt5 import QtCore, QtGui, QtWidgets
from PyQt5.QtCore import Qt, QThread, pyqtSignal
from PyQt5.QtWidgets import (
    QApplication, QMainWindow, QWidget, QFileDialog, QMessageBox,
    QVBoxLayout, QHBoxLayout, QLabel, QPushButton, QLineEdit, QCheckBox,
    QGroupBox, QTableWidget, QTableWidgetItem, QHeaderView, QTextEdit, QTabWidget
)
# Matplotlib Canvas
from matplotlib.backends.backend_qt5agg import FigureCanvasQTAgg as FigureCanvas
from matplotlib.figure import Figure

# Word export
from docx import Document
from docx.shared import Inches


# =========================
# ====== ANALYZER =========
# =========================

URL_GENERIC_RE = re.compile(
    r'(?:(?:hxxp|http|https):\/\/)?[a-zA-Z0-9\-\.\[\]]+\.[a-zA-Z]{2,}(?:[:\/][^\s<>"\)\]]*)?',
    re.IGNORECASE
)
IP_RE = re.compile(r'\b(?:\d{1,3}\.){3}\d{1,3}\b')
HASH_RE = re.compile(r'\b[a-fA-F0-9]{32}\b|\b[a-fA-F0-9]{40}\b|\b[a-fA-F0-9]{64}\b')

SHORTENERS = {"bit.ly","goo.gl","t.co","tinyurl.com","ow.ly","is.gd","buff.ly","cutt.ly","lnkd.in","rebrand.ly"}
SUS_TLDS = {"zip","mov","country","gq","work","click","top","xyz","support","help","pics"}

ATTACK_TECHNIQUES = {
    "T1566": "Phishing",
    "T1566.001": "Spearphishing Attachment",
    "T1566.002": "Spearphishing Link",
    "T1204": "User Execution",
    "T1140": "Deobfuscate/Decode",
    "T1059": "Command and Scripting Interpreter",
    "T1105": "Ingress Tool Transfer",
}

DEFAULT_CFG: Dict[str, Any] = {
    "apis": {
        "virustotal_api_key": "06e809ecf15d400ac95f5376765eff953e4225c0846d83a795fcb1b669d674c2",
        "otx_api_key": "ffa4f89b675c44f5f4d04b7b823720212095ef37c300c0564ffe81c78546df48",
        "urlscan_api_key": "0198e4cd-9856-71fe-a91d-e4f249908c2e",
        "greynoise_api_key": "fa909904-bf79-4cad-8f84-a43aa2659632",
        "abuseipdb_api_key": "fcab3a321d5debbab5b6f8fd32bf7847a9de868935f42daae14403a951c5490f3e6c7bd3ed630a3c"
    },
    "network": {
        "user_agent": "Mansimov-Phish-Analyzer/1.0",
        "timeout_sec": 10,
        "verify_tls": True,
        "safe_http_methods": ["HEAD", "GET"],
        "expand_max_redirects": 5
    },
    "scoring": {
        "weights": {
            "url_new_domain": 10,
            "url_sus_tld": 8,
            "url_ip_host": 12,
            "url_shortener": 6,
            "macro_autoexec": 20,
            "attachment_exe": 25,
            "attachment_script": 15,
            "vt_malicious_detection": 30,
            "brand_impersonation": 12,
            "dmarc_fail": 10,
            "replyto_mismatch": 8,
            "unicode_spoof": 10,
            "otx_pulse_hit": 10,
            "greynoise_malicious": 8,
            "abuseipdb_reports": 8
        }
    },
    "output": {
        "html_report": True,
        "markdown_report": True,
        "json_report": True
    },
    "yara": {
        "rules_path": "yara_rules"
    }
}

def load_config(path: Optional[str]) -> Dict[str, Any]:
    if not path:
        return DEFAULT_CFG
    p = Path(path)
    if not p.exists():
        return DEFAULT_CFG
    with open(p, "r", encoding="utf-8") as f:
        d = yaml.safe_load(f) or {}
    # Merge into defaults (shallow)
    cfg = json.loads(json.dumps(DEFAULT_CFG))
    for k, v in d.items():
        cfg[k] = v
    return cfg

def safe_decode(bytes_or_str: Any) -> str:
    if isinstance(bytes_or_str, str):
        return bytes_or_str
    if bytes_or_str is None:
        return ""
    guess = chardet.detect(bytes_or_str)
    enc = guess.get("encoding") or "utf-8"
    try:
        return bytes_or_str.decode(enc, errors="replace")
    except Exception:
        return bytes_or_str.decode("utf-8", errors="replace")

def normalize_url(u: str) -> str:
    s = u.strip()
    s = s.replace("[.]", ".").replace("[DOT]", ".").replace("[dot]", ".")
    s = s.replace("hxxp://", "http://").replace("hxxps://", "https://")
    s = re.sub(r'\s+', '', s)
    return s

def parse_eml(path: Path) -> Dict[str, Any]:
    with open(path, "rb") as f:
        msg = email.message_from_binary_file(f, policy=email.policy.default)
    return message_to_dict(msg)

def parse_msg(path: Path) -> Dict[str, Any]:
    if extract_msg is None:
        raise RuntimeError("extract_msg not installed")
    msg = extract_msg.Message(str(path))
    headers = {k: v for k, v in msg.headerItems()}
    dt = None
    try:
        dt = parsedate_to_datetime(headers.get("Date")) if headers.get("Date") else None
    except Exception:
        dt = None
    attachments = []
    for att in msg.attachments:
        attachments.append({
            "filename": att.longFilename or att.shortFilename,
            "content": att.data,
            "content_type": "application/octet-stream"
        })
    return {
        "headers": headers,
        "date": dt.isoformat() if dt else None,
        "subject": msg.subject or "",
        "from": headers.get("From", ""),
        "to": headers.get("To", ""),
        "reply_to": headers.get("Reply-To", ""),
        "body_text": safe_decode(msg.body.encode("utf-8")) if msg.body else "",
        "body_html": safe_decode((msg.htmlBody or "").encode("utf-8")),
        "attachments": attachments,
        "raw_source_path": str(path)
    }

def message_to_dict(msg: email.message.EmailMessage) -> Dict[str, Any]:
    headers = {k: str(make_header(decode_header(v))) for k, v in msg.items()}
    dt = None
    try:
        dt = parsedate_to_datetime(headers.get("Date")) if headers.get("Date") else None
    except Exception:
        dt = None
    body_text, body_html = "", ""
    attachments = []
    if msg.is_multipart():
        for part in msg.walk():
            ctype = part.get_content_type() or ""
            disp = (part.get("Content-Disposition") or "").lower()
            if ctype == "text/plain" and "attachment" not in disp:
                body_text += safe_decode(part.get_payload(decode=True))
            elif ctype == "text/html" and "attachment" not in disp:
                body_html += safe_decode(part.get_payload(decode=True))
            elif "attachment" in disp or part.get_filename():
                attachments.append({
                    "filename": part.get_filename(),
                    "content_type": ctype,
                    "content": part.get_payload(decode=True) or b""
                })
    else:
        ctype = msg.get_content_type()
        payload = msg.get_payload(decode=True) or b""
        if ctype == "text/plain":
            body_text = safe_decode(payload)
        elif ctype == "text/html":
            body_html = safe_decode(payload)
    return {
        "headers": headers,
        "date": dt.isoformat() if dt else None,
        "subject": headers.get("Subject", ""),
        "from": headers.get("From", ""),
        "to": headers.get("To", ""),
        "reply_to": headers.get("Reply-To", ""),
        "body_text": body_text,
        "body_html": body_html,
        "attachments": attachments,
        "raw_source_path": None
    }

def extract_urls(text: str) -> List[str]:
    urls = set()
    for m in URL_GENERIC_RE.finditer(text):
        urls.add(normalize_url(m.group(0)))
    return sorted(u for u in urls if "." in u)

def extract_iocs(email_obj: Dict[str, Any]) -> Dict[str, List[str]]:
    texts = " ".join([
        email_obj.get("subject",""),
        email_obj.get("from",""),
        email_obj.get("to",""),
        email_obj.get("reply_to",""),
        email_obj.get("body_text",""),
        email_obj.get("body_html","")
    ])
    urls = set(extract_urls(texts))
    ips = set(IP_RE.findall(texts))
    hashes = set(HASH_RE.findall(texts))
    soup = BeautifulSoup(email_obj.get("body_html",""), "lxml")
    for tag in soup.find_all(["a","img","form","script"]):
        for attr in ["href","src","action"]:
            val = tag.get(attr)
            if val:
                urls.add(normalize_url(val))
    domains = set()
    for u in list(urls):
        if "://" not in u:
            u = "http://" + u
        try:
            ext = tldextract.extract(u)
            if ext.domain and ext.suffix:
                domains.add(".".join(p for p in [ext.subdomain, ext.domain, ext.suffix] if p))
        except Exception:
            pass
    return {
        "urls": sorted(urls),
        "ips": sorted(ips),
        "domains": sorted(domains),
        "hashes": sorted(hashes)
    }

def header_auth_findings(headers: Dict[str,str]) -> List[str]:
    f = []
    spf = headers.get("Received-SPF","")
    authres = headers.get("Authentication-Results","")
    if "fail" in spf.lower():
        f.append("SPF fail observed in Received-SPF")
    if "dmarc=fail" in authres.lower():
        f.append("DMARC fail in Authentication-Results")
    if "dkim=fail" in authres.lower():
        f.append("DKIM fail in Authentication-Results")
    rp = headers.get("Return-Path","")
    frm = headers.get("From","")
    if rp and frm and rp.split("@")[-1].strip("<> ") != frm.split("@")[-1].strip("<> "):
        f.append("Return-Path domain mismatch with From")
    return f

def brand_impersonation_indicators(subject: str, body_html: str) -> List[str]:
    brands = ["Microsoft", "Office 365", "PayPal", "DocuSign", "Adobe", "Amazon", "OneDrive", "Dropbox", "Okta", "GitHub", "Google"]
    hits = []
    for b in brands:
        if re.search(rf'\b{re.escape(b)}\b', subject, re.I) or re.search(rf'\b{re.escape(b)}\b', body_html, re.I):
            hits.append(f"Brand keyword present: {b}")
    return hits

def contains_unicode_spoofers(s: str) -> bool:
    if re.search(r'[\u200B-\u200D\uFEFF]', s):
        return True
    return bool(re.search(r'[\u0400-\u04FF]', s))

def hash_bytes(b: bytes) -> Dict[str,str]:
    return {
        "md5": hashlib.md5(b).hexdigest(),
        "sha1": hashlib.sha1(b).hexdigest(),
        "sha256": hashlib.sha256(b).hexdigest()
    }

def analyze_attachments(atts: List[Dict[str,Any]], yara_rules: Optional[yara.Rules]) -> List[Dict[str,Any]]:
    results = []
    for att in atts:
        name = att.get("filename") or "unknown"
        ctype = att.get("content_type") or "application/octet-stream"
        content: bytes = att.get("content") or b""
        h = hash_bytes(content)
        finding = {
            "filename": name,
            "content_type": ctype,
            "size": len(content),
            "hashes": h,
            "macro_analysis": None,
            "pe_analysis": None,
            "yara_matches": [],
            "pe": False,
            "office_doc": False,
            "script_like": False
        }
        if yara_rules:
            try:
                m = yara_rules.match(data=content, timeout=3.0)
                finding["yara_matches"] = [str(i) for i in m]
            except Exception:
                pass
        if VBA_Parser is not None:
            try:
                vba = VBA_Parser(filename=name, data=content)
                finding["office_doc"] = vba.detect_vba_macros()
                if finding["office_doc"]:
                    autoexec = False
                    suspicious = []
                    for (_fname, _stream_path, _vba_filename, vba_code) in vba.extract_macros():
                        if not vba_code:
                            continue
                        if re.search(r'Auto(?:Open|_?Open|Close|_?Close|Workbook_Open|Document_Open)', vba_code, re.I):
                            autoexec = True
                        for pat in [r'CreateObject\("WScript\.Shell"\)', r'Shell\(', r'GetObject\(', r'Execute\(', r'Bitstransfer', r'URLDownloadToFile', r'WinHttp', r'MSXML2\.XMLHTTP']:
                            if re.search(pat, vba_code, re.I):
                                suspicious.append(pat)
                    finding["macro_analysis"] = {
                        "autoexec": autoexec,
                        "suspicious_patterns": list(sorted(set(suspicious)))
                    }
            except Exception:
                pass
        try:
            pe = pefile.PE(data=content, fast_load=True)
            finding["pe"] = True
            finding["pe_analysis"] = {
                "is_dll": bool(pe.FILE_HEADER.Characteristics & 0x2000),
                "machine": hex(pe.FILE_HEADER.Machine),
                "number_of_sections": pe.FILE_HEADER.NumberOfSections
            }
        except Exception:
            pass
        if name.lower().endswith((".vbs",".js",".jse",".wsf",".ps1",".cmd",".bat",".lnk",".hta",".vbe")):
            finding["script_like"] = True
        results.append(finding)
    return results

async def expand_url(session: aiohttp.ClientSession, url: str, max_redirects: int) -> Dict[str, Any]:
    url_norm = normalize_url(url)
    if not url_norm.startswith("http"):
        url_norm = "http://" + url_norm
    res = {"input": url, "expanded": url_norm, "chain": [], "status": None}
    try:
        async with session.get(url_norm, allow_redirects=False) as r:
            res["status"] = r.status
            res["chain"].append({"url": url_norm, "status": r.status})
            redirects = 0
            next_url = r.headers.get("Location")
            while next_url and redirects < max_redirects:
                if not next_url.startswith("http"):
                    from yarl import URL as YURL
                    next_url = str(YURL(url_norm).join(YURL(next_url)))
                res["chain"].append({"url": next_url, "status": None})
                async with session.get(next_url, allow_redirects=False) as nr:
                    res["chain"][-1]["status"] = nr.status
                    url_norm = next_url
                    next_url = nr.headers.get("Location")
                    redirects += 1
            res["expanded"] = url_norm
    except Exception as e:
        res["error"] = str(e)
    return res

async def enrich_iocs(cfg: Dict[str,Any], iocs: Dict[str, List[str]]) -> Dict[str,Any]:
    out = {"virustotal": {}, "otx": {}, "urlscan": {}, "greynoise": {}, "abuseipdb": {}, "whois": {}}
    headers = {"User-Agent": cfg["network"]["user_agent"]}
    timeout = aiohttp.ClientTimeout(total=cfg["network"]["timeout_sec"])
    ssl = cfg["network"]["verify_tls"]
    async with aiohttp.ClientSession(timeout=timeout, headers=headers) as session:
        vt_key = cfg["apis"].get("virustotal_api_key")
        if vt_key:
            for h in iocs["hashes"][:25]:
                url = f"https://www.virustotal.com/api/v3/files/{h}"
                try:
                    async with session.get(url, headers={"x-apikey": vt_key}, ssl=ssl) as r:
                        if r.status == 200:
                            out["virustotal"][h] = await r.json()
                except Exception:
                    pass
            for d in iocs["domains"][:25]:
                url = f"https://www.virustotal.com/api/v3/domains/{d}"
                try:
                    async with session.get(url, headers={"x-apikey": vt_key}, ssl=ssl) as r:
                        if r.status == 200:
                            out["virustotal"][d] = await r.json()
                except Exception:
                    pass
        us_key = cfg["apis"].get("urlscan_api_key")
        if us_key:
            for u in iocs["urls"][:25]:
                try:
                    q = f"https://urlscan.io/api/v1/search/?q=domain:{u}"
                    headers2 = {"API-Key": us_key, "Content-Type": "application/json"}
                    async with session.get(q, headers=headers2, ssl=ssl) as r:
                        if r.status == 200:
                            out["urlscan"][u] = await r.json()
                except Exception:
                    pass
        otx_key = cfg["apis"].get("otx_api_key")
        if otx_key:
            for d in iocs["domains"][:25]:
                try:
                    u = f"https://otx.alienvault.com/api/v1/indicators/domain/{d}/general"
                    async with session.get(u, headers={"X-OTX-API-KEY": otx_key}, ssl=ssl) as r:
                        if r.status == 200:
                            out["otx"][d] = await r.json()
                except Exception:
                    pass
        gn_key = cfg["apis"].get("greynoise_api_key")
        if gn_key:
            for ip in iocs["ips"][:50]:
                try:
                    u = f"https://api.greynoise.io/v3/community/{ip}"
                    async with session.get(u, headers={"key": gn_key}, ssl=ssl) as r:
                        if r.status == 200:
                            out["greynoise"][ip] = await r.json()
                except Exception:
                    pass
        ab_key = cfg["apis"].get("abuseipdb_api_key")
        if ab_key:
            for ip in iocs["ips"][:50]:
                try:
                    u = f"https://api.abuseipdb.com/api/v2/check?ipAddress={ip}&maxAgeInDays=90"
                    async with session.get(u, headers={"Key": ab_key, "Accept": "application/json"}, ssl=ssl) as r:
                        if r.status == 200:
                            out["abuseipdb"][ip] = await r.json()
                except Exception:
                    pass
        for d in iocs["domains"][:25]:
            try:
                w = whois.whois(d)
                out["whois"][d] = {
                    "creation_date": str(w.creation_date),
                    "updated_date": str(w.updated_date),
                    "registrar": w.registrar,
                    "name_servers": w.name_servers
                }
            except Exception:
                pass
    return out

def score_findings(cfg: Dict[str,Any], email_obj: Dict[str,Any], iocs: Dict[str,List[str]], att_results: List[Dict[str,Any]], enrich: Dict[str,Any]) -> Tuple[int, List[str], List[str]]:
    w = cfg["scoring"]["weights"]
    score = 0
    tags = set()
    notes = []

    for u in iocs["urls"]:
        try:
            ext = tldextract.extract(u if u.startswith("http") else "http://" + u)
            if ext and ext.suffix:
                if ext.suffix in SUS_TLDS:
                    score += w["url_sus_tld"]; notes.append(f"suspicious TLD: .{ext.suffix}")
        except Exception:
            pass
        if re.match(r'^https?://\d{1,3}(?:\.\d{1,3}){3}', u):
            score += w["url_ip_host"]; notes.append("URL host is raw IP")
        host = (tldextract.extract(u).fqdn or "").lower()
        if any(host.endswith(s) for s in SHORTENERS):
            score += w["url_shortener"]; notes.append("URL shortener used")
        tags.add("T1566.002")

    for a in att_results:
        if a["macro_analysis"]:
            if a["macro_analysis"]["autoexec"]:
                score += w["macro_autoexec"]; notes.append(f"Macro autoexec in {a['filename']}")
            if a["macro_analysis"]["suspicious_patterns"]:
                score += 8; notes.append(f"Suspicious VBA in {a['filename']}")
            tags.add("T1566.001"); tags.add("T1204")
        if a["pe"]:
            score += w["attachment_exe"]; notes.append(f"Executable attachment: {a['filename']}")
            tags.add("T1204")
        if a["script_like"]:
            score += w["attachment_script"]; notes.append(f"Script-like attachment: {a['filename']}")
        if a["yara_matches"]:
            score += 10; notes.append(f"YARA hits: {', '.join(a['yara_matches'])}")

    for h in header_auth_findings(email_obj["headers"]):
        notes.append(h)
    if any("DMARC fail" in n for n in notes):
        score += w["dmarc_fail"]
    if "Return-Path domain mismatch with From" in notes:
        score += w["replyto_mismatch"]

    brand_hits = brand_impersonation_indicators(email_obj.get("subject",""), email_obj.get("body_html",""))
    if brand_hits:
        notes.extend(brand_hits)
        score += w["brand_impersonation"]

    if contains_unicode_spoofers(email_obj.get("subject","") + email_obj.get("body_html","")):
        score += w["unicode_spoof"]; notes.append("Unicode zero-width/homoglyphs present")

    vt = enrich.get("virustotal", {})
    for key, val in vt.items():
        try:
            stats = val["data"]["attributes"]["last_analysis_stats"]
            if stats.get("malicious", 0) >= 3:
                score += w["vt_malicious_detection"]
                notes.append(f"VirusTotal indicates malicious for {key} ({stats.get('malicious')} vendors)")
        except Exception:
            pass

    for ip, v in enrich.get("abuseipdb", {}).items():
        try:
            rep = v["data"]["totalReports"]
            if rep and rep >= 5:
                score += w["abuseipdb_reports"]; notes.append(f"AbuseIPDB {ip} reports: {rep}")
        except Exception:
            pass

    for ip, v in enrich.get("greynoise", {}).items():
        try:
            if v.get("classification") == "malicious":
                score += w["greynoise_malicious"]; notes.append(f"GreyNoise malicious IP: {ip}")
        except Exception:
            pass

    for d, v in enrich.get("otx", {}).items():
        try:
            pulses = v.get("pulse_info", {}).get("count", 0)
            if pulses > 0:
                score += w["otx_pulse_hit"]; notes.append(f"OTX pulses hit: {d} ({pulses})")
        except Exception:
            pass

    if any("Macro" in n for n in notes) or any(a["macro_analysis"] for a in att_results):
        tags.add("T1140")
    if any(("URL shortener" in n) or ("URL host is raw IP" in n) for n in notes):
        tags.add("T1105")

    return score, sorted(tags), notes

def render_markdown(email_obj: Dict[str,Any], iocs: Dict[str,List[str]], att_results: List[Dict[str,Any]], score: int, tags: List[str], notes: List[str]) -> str:
    md = []
    md.append(f"# Phishing Analysis Report\n")
    md.append(f"- Subject: {email_obj.get('subject','')}")
    md.append(f"- From: {email_obj.get('from','')}")
    md.append(f"- To: {email_obj.get('to','')}")
    md.append(f"- Date: {email_obj.get('date')}")
    md.append(f"- Score: {score}")
    items = []
    for t in tags:
        technique = ATTACK_TECHNIQUES.get(t, "?")
        items.append(f"{t} ({technique})")
    
    line = "- MITRE ATT&CK: " + ", ".join(items)
    md.append(line)
    #md.append(f"- MITRE ATT&CK: {', '.join([f'{t} ({ATTACK_TECHNIQUES.get(t,\"?\")})' for t in tags])}")
    md.append("\n## Findings")
    for n in notes:
        md.append(f"- {n}")
    md.append("\n## IOCs")
    md.append(f"- Domains: {', '.join(iocs['domains'])}")
    md.append(f"- IPs: {', '.join(iocs['ips'])}")
    md.append(f"- URLs:")
    for u in iocs["urls"]:
        md.append(f"  - {u}")
    md.append(f"- Hashes: {', '.join(iocs['hashes'])}")
    md.append("\n## Attachments")
    for a in att_results:
        md.append(f"- {a['filename']} ({a['content_type']}, {a['size']} bytes)")
        md.append(f"  - Hashes: md5={a['hashes']['md5']} sha256={a['hashes']['sha256']}")
        if a["macro_analysis"]:
            md.append(f"  - Macro autoexec: {a['macro_analysis']['autoexec']}")
            if a["macro_analysis"]["suspicious_patterns"]:
                md.append(f"  - Suspicious VBA: {', '.join(a['macro_analysis']['suspicious_patterns'])}")
        if a["pe"]:
            md.append(f"  - PE: {json.dumps(a['pe_analysis'])}")
        if a["yara_matches"]:
            md.append(f"  - YARA: {', '.join(a['yara_matches'])}")
    return "\n".join(md)

async def analyze_email_async(
    email_path: str,
    cfg_path: Optional[str],
    yara_dir: Optional[str],
    expand: bool
) -> Dict[str, Any]:
    cfg = load_config(cfg_path)
    p = Path(email_path)
    if p.suffix.lower() == ".eml":
        email_obj = parse_eml(p)
    elif p.suffix.lower() == ".msg":
        email_obj = parse_msg(p)
    else:
        raise RuntimeError("Unsupported file type. Use .eml or .msg")

    yara_rules = None
    if yara_dir:
        yr = Path(yara_dir)
        if yr.exists():
            rule_files = [str(r) for r in yr.rglob("*.yar")]
            if rule_files:
                filemap = {f"r{i}": rf for i, rf in enumerate(rule_files)}
                try:
                    yara_rules = yara.compile(filepaths=filemap)
                except Exception as e:
                    # proceed without YARA
                    yara_rules = None

    att_results = analyze_attachments(email_obj["attachments"], yara_rules)
    iocs = extract_iocs(email_obj)

    expansions = []
    if expand and iocs["urls"]:
        timeout = aiohttp.ClientTimeout(total=cfg["network"]["timeout_sec"])
        async with aiohttp.ClientSession(timeout=timeout, headers={"User-Agent": cfg["network"]["user_agent"]}) as session:
            tasks = [expand_url(session, u, cfg["network"]["expand_max_redirects"]) for u in iocs["urls"][:30]]
            expansions = await asyncio.gather(*tasks)
        expanded = [e.get("expanded", e["input"]) for e in expansions]
        iocs["urls"] = sorted(set(expanded + iocs["urls"]))

    enrichment = await enrich_iocs(cfg, iocs)
    score, tags, notes = score_findings(cfg, email_obj, iocs, att_results, enrichment)
    md = render_markdown(email_obj, iocs, att_results, score, tags, notes)

    return {
        "email": email_obj,
        "iocs": iocs,
        "attachments": att_results,
        "score": score,
        "tags": tags,
        "notes": notes,
        "enrichment": enrichment,
        "expansions": expansions,
        "markdown": md,
        "config_used": cfg
    }


# =========================
# ====== WORKER ===========
# =========================

class AnalyzerWorker(QThread):
    finished = pyqtSignal(dict, str)  # result, error_message

    def __init__(self, email_path: str, cfg_path: Optional[str], yara_dir: Optional[str], expand: bool):
        super().__init__()
        self.email_path = email_path
        self.cfg_path = cfg_path
        self.yara_dir = yara_dir
        self.expand = expand

    def run(self):
        try:
            result = asyncio.run(analyze_email_async(self.email_path, self.cfg_path, self.yara_dir, self.expand))
            self.finished.emit(result, "")
        except Exception as e:
            self.finished.emit({}, str(e))


# =========================
# ====== CHARTS ===========
# =========================

class MplCanvas(FigureCanvas):
    def __init__(self, width=5, height=3, dpi=100):
        self.fig = Figure(figsize=(width, height), dpi=dpi)
        self.ax = self.fig.add_subplot(121)  # left plot
        self.ax2 = self.fig.add_subplot(122) # right plot
        super().__init__(self.fig)

    def draw_ioc_bar(self, ioc_counts: Dict[str, int]):
        self.ax.clear()
        keys = list(ioc_counts.keys())
        vals = [ioc_counts[k] for k in keys]
        self.ax.bar(keys, vals, color=["#2b8a3e","#1971c2","#e8590c","#862e9c"])
        self.ax.set_title("IOC Counts")
        self.ax.set_ylabel("Count")
        self.ax.set_xticklabels(keys, rotation=20, ha="right")
        self.fig.tight_layout()

    def draw_score_donut(self, score: int):
        # Clamp score to 0..100 for display purposes
        display_score = max(0, min(100, score))
        self.ax2.clear()
        sizes = [display_score, 100 - display_score]
        colors = ["#d9480f", "#e9ecef"] if display_score >= 60 else ("#f08c00" if display_score >= 30 else "#2f9e44", "#e9ecef")
        wedges, _ = self.ax2.pie(sizes, colors=colors, startangle=90, counterclock=False, wedgeprops=dict(width=0.35))
        self.ax2.set_aspect('equal')
        self.ax2.set_title(f"Score: {score}")
        self.fig.tight_layout()


# =========================
# ====== GUI MAIN =========
# =========================

class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Phishing Email Analyzer - Desktop")
        self.setMinimumSize(1100, 700)
        self.current_result: Optional[Dict[str, Any]] = None
        self.current_email_path: Optional[str] = None
        self.chart_canvas = MplCanvas(width=8, height=3.2, dpi=100)

        # Top controls
        top_box = QGroupBox("Inputs")
        self.email_edit = QLineEdit()
        self.email_btn = QPushButton("Browse Email (.eml/.msg)")
        self.cfg_edit = QLineEdit()
        self.cfg_btn = QPushButton("Config (config.yaml)")
        self.yara_edit = QLineEdit()
        self.yara_btn = QPushButton("YARA Rules Dir")
        self.expand_chk = QCheckBox("Expand URLs")
        self.analyze_btn = QPushButton("Analyze")
        self.export_btn = QPushButton("Export Word Report")
        self.export_btn.setEnabled(False)

        hl1 = QHBoxLayout()
        hl1.addWidget(QLabel("Email:"))
        hl1.addWidget(self.email_edit)
        hl1.addWidget(self.email_btn)

        hl2 = QHBoxLayout()
        hl2.addWidget(QLabel("Config:"))
        hl2.addWidget(self.cfg_edit)
        hl2.addWidget(self.cfg_btn)

        hl3 = QHBoxLayout()
        hl3.addWidget(QLabel("YARA:"))
        hl3.addWidget(self.yara_edit)
        hl3.addWidget(self.yara_btn)

        hl4 = QHBoxLayout()
        hl4.addWidget(self.expand_chk)
        hl4.addStretch()
        hl4.addWidget(self.analyze_btn)
        hl4.addWidget(self.export_btn)

        vtop = QVBoxLayout()
        vtop.addLayout(hl1)
        vtop.addLayout(hl2)
        vtop.addLayout(hl3)
        vtop.addLayout(hl4)
        top_box.setLayout(vtop)

        # Tabs
        self.tabs = QTabWidget()
        # Summary
        self.summary_text = QTextEdit(); self.summary_text.setReadOnly(True)
        tab_summary = QWidget(); v1 = QVBoxLayout(); v1.addWidget(self.summary_text); tab_summary.setLayout(v1)
        self.tabs.addTab(tab_summary, "Summary")

        # IOCs
        self.ioc_table = QTableWidget(0, 4)
        self.ioc_table.setHorizontalHeaderLabels(["Domains", "IPs", "URLs", "Hashes"])
        self.ioc_table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        tab_ioc = QWidget(); v2 = QVBoxLayout(); v2.addWidget(self.ioc_table); tab_ioc.setLayout(v2)
        self.tabs.addTab(tab_ioc, "IOCs")

        # Attachments
        self.att_table = QTableWidget(0, 7)
        self.att_table.setHorizontalHeaderLabels(["Filename","Type","Size","MD5","SHA256","Office/Macro","YARA/PE/Script"])
        self.att_table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        tab_att = QWidget(); v3 = QVBoxLayout(); v3.addWidget(self.att_table); tab_att.setLayout(v3)
        self.tabs.addTab(tab_att, "Attachments")

        # Enrichment
        self.enrich_text = QTextEdit(); self.enrich_text.setReadOnly(True)
        tab_enr = QWidget(); v4 = QVBoxLayout(); v4.addWidget(self.enrich_text); tab_enr.setLayout(v4)
        self.tabs.addTab(tab_enr, "Enrichment")

        # Charts
        tab_charts = QWidget(); v5 = QVBoxLayout(); v5.addWidget(self.chart_canvas); tab_charts.setLayout(v5)
        self.tabs.addTab(tab_charts, "Dashboard")

        # Central layout
        central = QWidget()
        layout = QVBoxLayout()
        layout.addWidget(top_box)
        layout.addWidget(self.tabs)
        central.setLayout(layout)
        self.setCentralWidget(central)

        # Signals
        self.email_btn.clicked.connect(self.pick_email)
        self.cfg_btn.clicked.connect(self.pick_config)
        self.yara_btn.clicked.connect(self.pick_yara)
        self.analyze_btn.clicked.connect(self.start_analyze)
        self.export_btn.clicked.connect(self.export_docx)

        # Style
        self.statusBar().showMessage("Ready")

    # ---------- File pickers ----------
    def pick_email(self):
        path, _ = QFileDialog.getOpenFileName(self, "Select Email File", "", "Email Files (*.eml *.msg)")
        if path:
            self.email_edit.setText(path)

    def pick_config(self):
        path, _ = QFileDialog.getOpenFileName(self, "Select Config YAML", "", "YAML (*.yaml *.yml)")
        if path:
            self.cfg_edit.setText(path)

    def pick_yara(self):
        path = QFileDialog.getExistingDirectory(self, "Select YARA Rules Directory", "")
        if path:
            self.yara_edit.setText(path)

    # ---------- Analysis ----------
    def start_analyze(self):
        email_path = self.email_edit.text().strip()
        if not email_path or not Path(email_path).exists():
            QMessageBox.warning(self, "Missing", "Select a valid .eml or .msg file.")
            return
        cfg_path = self.cfg_edit.text().strip() or None
        yara_dir = self.yara_edit.text().strip() or None
        expand = self.expand_chk.isChecked()

        self.analyze_btn.setEnabled(False)
        self.export_btn.setEnabled(False)
        self.statusBar().showMessage("Analyzing...")
        self.summary_text.setPlainText("Running analysis...")
        self.worker = AnalyzerWorker(email_path, cfg_path, yara_dir, expand)
        self.worker.finished.connect(self.analysis_done)
        self.worker.start()

    def analysis_done(self, result: Dict[str, Any], err: str):
        self.analyze_btn.setEnabled(True)
        if err:
            QMessageBox.critical(self, "Analysis Failed", err)
            self.statusBar().showMessage("Error")
            return
        self.current_result = result
        self.current_email_path = self.email_edit.text().strip()
        self.populate_views(result)
        self.export_btn.setEnabled(True)
        self.statusBar().showMessage("Analysis complete")

    # ---------- Populate UI ----------
    def populate_views(self, res: Dict[str, Any]):
        email_obj = res["email"]
        iocs = res["iocs"]
        atts = res["attachments"]
        score = res["score"]
        tags = res["tags"]
        notes = res["notes"]
        enrichment = res["enrichment"]

        # Summary
        summary_lines = []
        summary_lines.append(f"Subject: {email_obj.get('subject','')}")
        summary_lines.append(f"From: {email_obj.get('from','')}")
        summary_lines.append(f"To: {email_obj.get('to','')}")
        summary_lines.append(f"Date: {email_obj.get('date')}")
        summary_lines.append(f"Score: {score}")
        items = []
        for t in tags:
            technique = ATTACK_TECHNIQUES.get(t, "?")
            items.append(f"{t} ({technique})")
        
        summary_line = "MITRE ATT&CK: " + ", ".join(items)
        summary_lines.append(summary_line)
        # summary_lines.append("MITRE ATT&CK: " + ", ".join([f"{t} ({ATTACK_TECHNIQUES.get(t,'?')})" for t in tags]))
        summary_lines.append("\nFindings:")
        for n in notes:
            summary_lines.append(f"- {n}")
        self.summary_text.setPlainText("\n".join(summary_lines))

        # IOC table
        rows = max(len(iocs["domains"]), len(iocs["ips"]), len(iocs["urls"]), len(iocs["hashes"]))
        self.ioc_table.setRowCount(rows if rows>0 else 1)
        for r in range(rows):
            self.ioc_table.setItem(r, 0, QTableWidgetItem(iocs["domains"][r] if r < len(iocs["domains"]) else ""))
            self.ioc_table.setItem(r, 1, QTableWidgetItem(iocs["ips"][r] if r < len(iocs["ips"]) else ""))
            self.ioc_table.setItem(r, 2, QTableWidgetItem(iocs["urls"][r] if r < len(iocs["urls"]) else ""))
            self.ioc_table.setItem(r, 3, QTableWidgetItem(iocs["hashes"][r] if r < len(iocs["hashes"]) else ""))

        # Attachments table
        self.att_table.setRowCount(len(atts) if atts else 1)
        if atts:
            for i, a in enumerate(atts):
                self.att_table.setItem(i, 0, QTableWidgetItem(a["filename"]))
                self.att_table.setItem(i, 1, QTableWidgetItem(a["content_type"]))
                self.att_table.setItem(i, 2, QTableWidgetItem(str(a["size"])))
                self.att_table.setItem(i, 3, QTableWidgetItem(a["hashes"]["md5"]))
                self.att_table.setItem(i, 4, QTableWidgetItem(a["hashes"]["sha256"]))
                macro_str = "No"
                if a["macro_analysis"]:
                    ma = a["macro_analysis"]
                    macro_str = f"AutoExec={ma['autoexec']}; Susp={len(ma['suspicious_patterns'])}"
                self.att_table.setItem(i, 5, QTableWidgetItem(("OfficeDoc " if a["office_doc"] else "") + macro_str))
                flags = []
                if a["yara_matches"]:
                    flags.append("YARA:" + ",".join(a["yara_matches"]))
                if a["pe"]:
                    flags.append("PE")
                if a["script_like"]:
                    flags.append("Script")
                self.att_table.setItem(i, 6, QTableWidgetItem(", ".join(flags)))
        else:
            for c in range(7):
                self.att_table.setItem(0, c, QTableWidgetItem(""))

        # Enrichment view (pretty JSON but trimmed)
        enr_slim = {
            "virustotal_keys": list(enrichment.get("virustotal", {}).keys())[:10],
            "otx_domains": list(enrichment.get("otx", {}).keys())[:10],
            "urlscan_urls": list(enrichment.get("urlscan", {}).keys())[:10],
            "greynoise_ips": list(enrichment.get("greynoise", {}).keys())[:10],
            "abuseipdb_ips": list(enrichment.get("abuseipdb", {}).keys())[:10],
            "whois_domains": list(enrichment.get("whois", {}).keys())[:10],
        }
        pretty = json.dumps(enr_slim, indent=2)
        self.enrich_text.setPlainText(pretty)

        # Charts
        ioc_counts = {
            "Domains": len(iocs["domains"]),
            "IPs": len(iocs["ips"]),
            "URLs": len(iocs["urls"]),
            "Hashes": len(iocs["hashes"]),
        }
        self.chart_canvas.draw_ioc_bar(ioc_counts)
        self.chart_canvas.draw_score_donut(score)
        self.chart_canvas.draw()

    # ---------- Export DOCX ----------
    def export_docx(self):
        if not self.current_result:
            QMessageBox.information(self, "No Data", "Run an analysis first.")
            return
        default_name = f"phish_report_{Path(self.current_email_path).stem}_{datetime.utcnow().strftime('%Y%m%dT%H%M%SZ')}.docx"
        out_path, _ = QFileDialog.getSaveFileName(self, "Save Report", default_name, "Word Document (*.docx)")
        if not out_path:
            return

        res = self.current_result
        email_obj = res["email"]
        iocs = res["iocs"]
        atts = res["attachments"]
        score = res["score"]
        tags = res["tags"]
        notes = res["notes"]
        md = res["markdown"]

        doc = Document()
        doc.add_heading("Phishing Analysis Report", 0)
        doc.add_paragraph(f"Generated: {datetime.utcnow().isoformat()}Z")
        doc.add_paragraph(f"Source file: {self.current_email_path}")

        doc.add_heading("Summary", level=1)
        p = doc.add_paragraph()
        p.add_run("Subject: ").bold = True; p.add_run(email_obj.get("subject",""))
        p = doc.add_paragraph()
        p.add_run("From: ").bold = True; p.add_run(email_obj.get("from",""))
        p = doc.add_paragraph()
        p.add_run("To: ").bold = True; p.add_run(email_obj.get("to",""))
        p = doc.add_paragraph()
        p.add_run("Date: ").bold = True; p.add_run(str(email_obj.get("date")))
        p = doc.add_paragraph()
        p.add_run("Score: ").bold = True; p.add_run(str(score))
        p = doc.add_paragraph()
        p.add_run("MITRE ATT&CK: ").bold = True; p.add_run(", ".join([f"{t} ({ATTACK_TECHNIQUES.get(t,'?')})" for t in tags]))

        doc.add_heading("Findings", level=1)
        for n in notes:
            doc.add_paragraph(n, style="List Bullet")

        doc.add_heading("IOCs", level=1)
        table = doc.add_table(rows=1, cols=4)
        hdr = table.rows[0].cells
        hdr[0].text = "Domains"; hdr[1].text = "IPs"; hdr[2].text = "URLs"; hdr[3].text = "Hashes"
        rows = max(len(iocs["domains"]), len(iocs["ips"]), len(iocs["urls"]), len(iocs["hashes"]))
        for r in range(rows):
            row_cells = table.add_row().cells
            row_cells[0].text = iocs["domains"][r] if r < len(iocs["domains"]) else ""
            row_cells[1].text = iocs["ips"][r] if r < len(iocs["ips"]) else ""
            row_cells[2].text = iocs["urls"][r] if r < len(iocs["urls"]) else ""
            row_cells[3].text = iocs["hashes"][r] if r < len(iocs["hashes"]) else ""

        doc.add_heading("Attachments", level=1)
        table2 = doc.add_table(rows=1, cols=7)
        hdr2 = table2.rows[0].cells
        hdr2[0].text="Filename"; hdr2[1].text="Type"; hdr2[2].text="Size"; hdr2[3].text="MD5"; hdr2[4].text="SHA256"; hdr2[5].text="Office/Macro"; hdr2[6].text="YARA/PE/Script"
        for a in atts:
            row = table2.add_row().cells
            row[0].text = a["filename"]
            row[1].text = a["content_type"]
            row[2].text = str(a["size"])
            row[3].text = a["hashes"]["md5"]
            row[4].text = a["hashes"]["sha256"]
            macro_str = "No"
            if a["macro_analysis"]:
                ma = a["macro_analysis"]
                macro_str = f"AutoExec={ma['autoexec']}; Susp={len(ma['suspicious_patterns'])}"
            row[5].text = ("OfficeDoc " if a["office_doc"] else "") + macro_str
            flags = []
            if a["yara_matches"]:
                flags.append("YARA:" + ",".join(a["yara_matches"]))
            if a["pe"]:
                flags.append("PE")
            if a["script_like"]:
                flags.append("Script")
            row[6].text = ", ".join(flags)

        # Add charts: save current canvas to temp PNGs
        tmp_dir = Path(os.getcwd()) / "_tmp_analyzer_images"
        tmp_dir.mkdir(exist_ok=True)
        chart_path = str(tmp_dir / "chart.png")
        self.chart_canvas.fig.savefig(chart_path, dpi=150, bbox_inches="tight")
        doc.add_heading("Dashboard", level=1)
        doc.add_picture(chart_path, width=Inches(6.5))

        # Add raw markdown (optional)
        doc.add_heading("Markdown Summary", level=1)
        doc.add_paragraph(md)

        try:
            doc.save(out_path)
            QMessageBox.information(self, "Saved", f"Report saved to:\n{out_path}")
        except Exception as e:
            QMessageBox.critical(self, "Save Failed", str(e))

# --- main ---
def main():
    app = QApplication(sys.argv)
    win = MainWindow()
    win.show()
    sys.exit(app.exec_())

if __name__ == "__main__":
    main()
